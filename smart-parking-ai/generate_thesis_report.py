from __future__ import annotations

import argparse
import csv
import json
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate thesis-ready ablation report from existing JSON artifacts")
    parser.add_argument(
        "--convert-summary",
        action="append",
        default=[],
        help="Conversion summary JSON path, repeatable (e.g. 2019 + 2020)",
    )
    parser.add_argument("--split-summary", required=True, help="Path to split_summary.json")
    parser.add_argument(
        "--train-reports-glob",
        default="artifacts/training/ccpd_mixed_yolo11n_ab*.json",
        help="Glob for training report JSON files",
    )
    parser.add_argument(
        "--eval-reports-glob",
        default="artifacts/eval/ccpd_mixed_yolo11n_ab*_test_eval.json",
        help="Glob for evaluation report JSON files",
    )
    parser.add_argument(
        "--output-docx",
        default="artifacts/paper/CCPD_mixed_ablation_report.docx",
        help="Word output path",
    )
    parser.add_argument(
        "--output-json",
        default="artifacts/paper/ablation_summary.json",
        help="Structured summary JSON output path",
    )
    parser.add_argument(
        "--output-csv",
        default="artifacts/paper/ablation_summary.csv",
        help="Tabular ablation CSV output path",
    )
    parser.add_argument(
        "--title",
        default="CCPD2019+CCPD2020 混合训练车牌检测消融实验报告",
        help="Document title",
    )
    return parser.parse_args()


def load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def infer_experiment_name(eval_report_path: Path, eval_payload: dict[str, Any]) -> str:
    stem = eval_report_path.stem
    for suffix in ("_test_eval", "_eval", "_test"):
        if stem.endswith(suffix):
            return stem[: -len(suffix)]
    tags = eval_payload.get("experiment", {}).get("tags", [])
    for tag in tags:
        if isinstance(tag, str) and "ab" in tag and "yolo11" in tag:
            return tag
    return stem


def safe_float(value: Any, default: float = 0.0) -> float:
    try:
        return float(value)
    except Exception:
        return float(default)


def yes_no(flag: bool) -> str:
    return "是" if flag else "否"


def fmt4(value: float) -> str:
    return f"{value:.4f}"


def fmt3(value: float) -> str:
    return f"{value:.3f}"


def build_rows(train_reports: list[Path], eval_reports: list[Path]) -> list[dict[str, Any]]:
    train_by_name: dict[str, dict[str, Any]] = {}
    for path in train_reports:
        payload = load_json(path)
        exp_name = payload.get("training", {}).get("name") or path.stem
        train_by_name[exp_name] = payload

    rows: list[dict[str, Any]] = []
    for eval_path in eval_reports:
        eval_payload = load_json(eval_path)
        exp_name = infer_experiment_name(eval_path, eval_payload)
        train_payload = train_by_name.get(exp_name, {})
        training = train_payload.get("training", {})
        extra_args = training.get("extra_train_args", {})
        experiment = train_payload.get("experiment", {})
        metrics = eval_payload.get("metrics", {})
        speed = metrics.get("speed", {}) if isinstance(metrics.get("speed"), dict) else {}
        row = {
            "experiment": exp_name,
            "imgsz": int(training.get("imgsz", eval_payload.get("imgsz", 0)) or 0),
            "epochs": int(training.get("epochs", 0) or 0),
            "batch": int(training.get("batch", eval_payload.get("batch", 0)) or 0),
            "fraction": safe_float(extra_args.get("fraction", 1.0), default=1.0),
            "cos_lr": bool(extra_args.get("cos_lr", False)),
            "precision": safe_float(metrics.get("box_mp", 0.0)),
            "recall": safe_float(metrics.get("box_mr", 0.0)),
            "map50": safe_float(metrics.get("box_map50", 0.0)),
            "map50_95": safe_float(metrics.get("box_map", 0.0)),
            "inference_ms": safe_float(speed.get("inference", 0.0)),
            "notes": str(experiment.get("notes", "")),
            "best_weights": str(training.get("best_weights", eval_payload.get("model", ""))),
            "run_dir": str(training.get("run_dir", "")),
            "train_report": str((Path(train_payload.get("training", {}).get("run_dir", "")) if training.get("run_dir") else "").as_posix())
            if False
            else "",
            "eval_report": str(eval_path),
        }
        rows.append(row)

    rows.sort(key=lambda item: item["experiment"])
    if not rows:
        raise RuntimeError("No ablation rows were built from report files")

    baseline = None
    for row in rows:
        if "ab0" in row["experiment"]:
            baseline = row
            break
    if baseline is None:
        baseline = rows[0]

    baseline_map = baseline["map50_95"]
    baseline_map50 = baseline["map50"]
    baseline_recall = baseline["recall"]
    for row in rows:
        row["delta_map50_95"] = row["map50_95"] - baseline_map
        row["delta_map50"] = row["map50"] - baseline_map50
        row["delta_recall"] = row["recall"] - baseline_recall
        row["is_baseline"] = row["experiment"] == baseline["experiment"]

    return rows


def write_csv(rows: list[dict[str, Any]], output_csv: Path) -> None:
    output_csv.parent.mkdir(parents=True, exist_ok=True)
    headers = [
        "experiment",
        "imgsz",
        "epochs",
        "batch",
        "fraction",
        "cos_lr",
        "precision",
        "recall",
        "map50",
        "map50_95",
        "delta_map50_95_vs_baseline",
        "delta_map50_vs_baseline",
        "delta_recall_vs_baseline",
        "inference_ms",
        "notes",
        "best_weights",
        "eval_report",
    ]
    with output_csv.open("w", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=headers)
        writer.writeheader()
        for row in rows:
            writer.writerow(
                {
                    "experiment": row["experiment"],
                    "imgsz": row["imgsz"],
                    "epochs": row["epochs"],
                    "batch": row["batch"],
                    "fraction": row["fraction"],
                    "cos_lr": row["cos_lr"],
                    "precision": row["precision"],
                    "recall": row["recall"],
                    "map50": row["map50"],
                    "map50_95": row["map50_95"],
                    "delta_map50_95_vs_baseline": row["delta_map50_95"],
                    "delta_map50_vs_baseline": row["delta_map50"],
                    "delta_recall_vs_baseline": row["delta_recall"],
                    "inference_ms": row["inference_ms"],
                    "notes": row["notes"],
                    "best_weights": row["best_weights"],
                    "eval_report": row["eval_report"],
                }
            )


def add_dataset_section(doc: Document, convert_summaries: list[dict[str, Any]], split_summary: dict[str, Any]) -> None:
    doc.add_heading("1. 数据集构建与划分", level=1)
    doc.add_paragraph("本实验将 CCPD2019 与 CCPD2020 合并后训练单类别车牌检测模型，类别为 plate。")

    conv_table = doc.add_table(rows=1, cols=5)
    conv_header = conv_table.rows[0].cells
    conv_header[0].text = "数据源"
    conv_header[1].text = "输入目录"
    conv_header[2].text = "扫描图像"
    conv_header[3].text = "转换成功"
    conv_header[4].text = "异常跳过"

    for payload in convert_summaries:
        row = conv_table.add_row().cells
        stats = payload.get("stats", {})
        data_source = Path(str(payload.get("input_dir", ""))).name or "unknown"
        skipped = (
            int(stats.get("skipped_parse_bbox", 0))
            + int(stats.get("skipped_read_image", 0))
            + int(stats.get("skipped_invalid_bbox", 0))
            + int(stats.get("skipped_small_box", 0))
            + int(stats.get("skipped_small_area_ratio", 0))
            + int(stats.get("skipped_write_error", 0))
        )
        row[0].text = data_source
        row[1].text = str(payload.get("input_dir", ""))
        row[2].text = str(stats.get("scanned_images", 0))
        row[3].text = str(stats.get("converted", 0))
        row[4].text = str(skipped)

    doc.add_paragraph("")
    split_table = doc.add_table(rows=2, cols=4)
    split_table.rows[0].cells[0].text = "总样本数"
    split_table.rows[0].cells[1].text = "训练集"
    split_table.rows[0].cells[2].text = "验证集"
    split_table.rows[0].cells[3].text = "测试集"
    split_table.rows[1].cells[0].text = str(split_summary.get("total_samples", 0))
    split_table.rows[1].cells[1].text = str(split_summary.get("splits", {}).get("train", 0))
    split_table.rows[1].cells[2].text = str(split_summary.get("splits", {}).get("val", 0))
    split_table.rows[1].cells[3].text = str(split_summary.get("splits", {}).get("test", 0))

    ratios = split_summary.get("ratios", {})
    doc.add_paragraph(
        "数据划分比例："
        f"train={ratios.get('train', 'N/A')} / val={ratios.get('val', 'N/A')} / test={ratios.get('test', 'N/A')}。"
    )


def add_ablation_design_section(doc: Document, rows: list[dict[str, Any]]) -> None:
    doc.add_heading("2. 消融实验设计", level=1)
    doc.add_paragraph("固定条件：YOLO11n，单卡 GPU，batch=8，epochs=5，fraction=0.2；对比输入尺寸与学习率调度策略。")

    table = doc.add_table(rows=1, cols=7)
    header = table.rows[0].cells
    header[0].text = "实验名"
    header[1].text = "输入尺寸"
    header[2].text = "cos_lr"
    header[3].text = "epochs"
    header[4].text = "fraction"
    header[5].text = "batch"
    header[6].text = "备注"
    for row_data in rows:
        row = table.add_row().cells
        row[0].text = row_data["experiment"]
        row[1].text = str(row_data["imgsz"])
        row[2].text = yes_no(row_data["cos_lr"])
        row[3].text = str(row_data["epochs"])
        row[4].text = str(row_data["fraction"])
        row[5].text = str(row_data["batch"])
        row[6].text = row_data["notes"] or "-"


def add_results_section(doc: Document, rows: list[dict[str, Any]]) -> dict[str, Any]:
    doc.add_heading("3. 消融实验结果（Test 集）", level=1)
    table = doc.add_table(rows=1, cols=8)
    header = table.rows[0].cells
    header[0].text = "实验名"
    header[1].text = "Precision"
    header[2].text = "Recall"
    header[3].text = "mAP@0.5"
    header[4].text = "mAP@0.5:0.95"
    header[5].text = "ΔmAP@0.5:0.95"
    header[6].text = "ΔRecall"
    header[7].text = "Inference(ms)"

    best = max(rows, key=lambda item: item["map50_95"])
    baseline = next((item for item in rows if item["is_baseline"]), rows[0])
    for row_data in rows:
        row = table.add_row().cells
        row[0].text = row_data["experiment"]
        row[1].text = fmt4(row_data["precision"])
        row[2].text = fmt4(row_data["recall"])
        row[3].text = fmt4(row_data["map50"])
        row[4].text = fmt4(row_data["map50_95"])
        row[5].text = f"{row_data['delta_map50_95']:+.4f}"
        row[6].text = f"{row_data['delta_recall']:+.4f}"
        row[7].text = fmt3(row_data["inference_ms"])

    doc.add_paragraph(
        "最佳实验："
        f"{best['experiment']}（mAP@0.5:0.95={fmt4(best['map50_95'])}，mAP@0.5={fmt4(best['map50'])}，Recall={fmt4(best['recall'])}）。"
    )
    doc.add_paragraph(
        "相对基线："
        f"ΔmAP@0.5:0.95={best['delta_map50_95']:+.4f}，"
        f"ΔmAP@0.5={best['delta_map50']:+.4f}，"
        f"ΔRecall={best['delta_recall']:+.4f}。"
    )
    return {"best": best, "baseline": baseline}


def add_figure_and_repro_section(doc: Document, rows: list[dict[str, Any]], split_summary_path: Path) -> None:
    doc.add_heading("4. 图示文件与复现实验命令", level=1)
    doc.add_paragraph("以下路径可直接用于论文插图：results.png、PR 曲线、混淆矩阵等。")
    for row in rows:
        run_dir = Path(row["run_dir"]) if row["run_dir"] else None
        if not run_dir:
            continue
        results_png = run_dir / "results.png"
        pr_curve = run_dir / "BoxPR_curve.png"
        conf_matrix = run_dir / "confusion_matrix.png"
        doc.add_paragraph(
            f"{row['experiment']}:\n"
            f"  - {results_png}\n"
            f"  - {pr_curve}\n"
            f"  - {conf_matrix}"
        )

    split_summary = load_json(split_summary_path)
    source_images = split_summary.get("source_images_dir", "")
    source_labels = split_summary.get("source_labels_dir", "")
    output_dir = split_summary.get("output_dir", "")
    dataset_yaml = split_summary.get("dataset_yaml", "")
    doc.add_paragraph("复现命令示例（使用 D:\\Python\\python.exe）：")
    doc.add_paragraph(
        "1) 数据划分：\n"
        f"D:\\Python\\python.exe prepare_dataset.py --images-dir {source_images} --labels-dir {source_labels} "
        f"--output-dir {output_dir} --train-ratio 0.7 --val-ratio 0.2 --test-ratio 0.1 --seed 42 --class-names plate"
    )
    doc.add_paragraph(
        "2) 基线训练：\n"
        "D:\\Python\\python.exe train_yolo.py --dataset-yaml "
        f"{dataset_yaml} --base-model D:\\project\\smart-parking-ai\\models\\yolo11n.pt --epochs 5 --imgsz 640 --batch 8 --device 0 "
        "--project runs/train --name ccpd_mixed_yolo11n_ab0_base_e5_f20 --train-arg fraction=0.2 "
        "--tag ccpd_mixed --tag ablation --tag baseline "
        "--save-report artifacts/training/ccpd_mixed_yolo11n_ab0_base_e5_f20.json"
    )
    doc.add_paragraph(
        "3) 改进训练（示例）：\n"
        "D:\\Python\\python.exe train_yolo.py --dataset-yaml "
        f"{dataset_yaml} --base-model D:\\project\\smart-parking-ai\\models\\yolo11n.pt --epochs 5 --imgsz 640 --batch 8 --device 0 "
        "--project runs/train --name ccpd_mixed_yolo11n_ab1_coslr_e5_f20 --train-arg fraction=0.2 --train-arg cos_lr=true "
        "--tag ccpd_mixed --tag ablation --tag coslr "
        "--save-report artifacts/training/ccpd_mixed_yolo11n_ab1_coslr_e5_f20.json"
    )
    doc.add_paragraph(
        "4) Test 评估：\n"
        "D:\\Python\\python.exe evaluate_yolo.py --model "
        "D:\\project\\smart-parking-ai\\runs\\detect\\runs\\train\\ccpd_mixed_yolo11n_ab1_coslr_e5_f20\\weights\\best.pt "
        f"--dataset-yaml {dataset_yaml} --split test --imgsz 640 --batch 8 --device 0 "
        "--save-report artifacts/eval/ccpd_mixed_yolo11n_ab1_coslr_e5_f20_test_eval.json"
    )


def add_limitation_section(doc: Document) -> None:
    doc.add_heading("5. 当前结论与后续建议", level=1)
    doc.add_paragraph("当前消融为快速对比设置（epochs=5, fraction=0.2），用于筛选方向，尚非最终收敛精度。")
    doc.add_paragraph("建议下一步将最佳配置扩展到完整训练（fraction=1.0，epochs>=50）并补充多随机种子复现实验。")
    doc.add_paragraph("若用于论文终稿，请补充难样本分析（夜间、遮挡、倾斜）与误检/漏检可视化统计。")


def main() -> int:
    args = parse_args()
    split_summary_path = Path(args.split_summary).expanduser().resolve()
    if not split_summary_path.is_file():
        raise RuntimeError(f"split summary not found: {split_summary_path}")

    convert_paths = [Path(item).expanduser().resolve() for item in args.convert_summary]
    for path in convert_paths:
        if not path.is_file():
            raise RuntimeError(f"convert summary not found: {path}")
    convert_payloads = [load_json(path) for path in convert_paths]

    train_reports = sorted(Path(".").glob(args.train_reports_glob))
    eval_reports = sorted(Path(".").glob(args.eval_reports_glob))
    if not train_reports:
        raise RuntimeError(f"no training reports matched: {args.train_reports_glob}")
    if not eval_reports:
        raise RuntimeError(f"no eval reports matched: {args.eval_reports_glob}")

    rows = build_rows(train_reports=train_reports, eval_reports=eval_reports)
    outcome = add_outcome(rows)

    output_json = Path(args.output_json).expanduser().resolve()
    output_csv = Path(args.output_csv).expanduser().resolve()
    output_docx = Path(args.output_docx).expanduser().resolve()
    output_json.parent.mkdir(parents=True, exist_ok=True)
    output_csv.parent.mkdir(parents=True, exist_ok=True)
    output_docx.parent.mkdir(parents=True, exist_ok=True)

    summary_payload = {
        "created_at": datetime.now().isoformat(timespec="seconds"),
        "title": args.title,
        "split_summary_path": str(split_summary_path),
        "convert_summaries": [str(path) for path in convert_paths],
        "experiments": rows,
        "best_experiment": outcome["best"],
        "baseline_experiment": outcome["baseline"],
    }
    output_json.write_text(json.dumps(summary_payload, ensure_ascii=False, indent=2), encoding="utf-8")
    write_csv(rows, output_csv)

    split_payload = load_json(split_summary_path)
    doc = Document()
    doc.add_heading(args.title, level=0)
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("本报告由脚本自动汇总训练与评估结果，可直接用于毕业设计实验章节初稿。")
    add_dataset_section(doc, convert_payloads, split_payload)
    add_ablation_design_section(doc, rows)
    add_results_section(doc, rows)
    add_figure_and_repro_section(doc, rows, split_summary_path)
    add_limitation_section(doc)
    doc.save(str(output_docx))

    print(json.dumps({"output_docx": str(output_docx), "output_csv": str(output_csv), "output_json": str(output_json)}, ensure_ascii=False, indent=2))
    return 0


def add_outcome(rows: list[dict[str, Any]]) -> dict[str, Any]:
    best = max(rows, key=lambda item: item["map50_95"])
    baseline = next((item for item in rows if item["is_baseline"]), rows[0])
    return {"best": best, "baseline": baseline}


if __name__ == "__main__":
    raise SystemExit(main())
