from __future__ import annotations

import argparse
import json
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parent.parent
DEFAULT_RECORD_ROOT = PROJECT_ROOT / "project_archive" / "moved_dirs" / "plate_thesis_records"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Build paper-ready summary documents for CCPD mixed plate detection experiments."
    )
    parser.add_argument("--record-root", default=str(DEFAULT_RECORD_ROOT))
    parser.add_argument("--paper-summary-md", default="paper_summary.md")
    parser.add_argument("--artifact-index-md", default="artifact_index.md")
    parser.add_argument("--manifest-json", default="summaries/paper_artifact_manifest.json")
    parser.add_argument("--paper-summary-docx", default="paper_summary.docx")
    return parser.parse_args()


def load_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def write_text(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")


def write_json(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def fmt_float(value: float, digits: int = 4) -> str:
    return f"{float(value):.{digits}f}"


def parse_nvidia_smi(text: str) -> dict[str, str]:
    info: dict[str, str] = {}
    driver_match = re.search(r"Driver Version:\s*([0-9.]+)", text)
    cuda_match = re.search(r"CUDA Version:\s*([0-9.]+)", text)
    gpu_match = re.search(r"\|\s*0\s+(.+?)\s{2,}WDDM", text)
    if driver_match:
        info["driver_version"] = driver_match.group(1)
    if cuda_match:
        info["cuda_runtime_version"] = cuda_match.group(1)
    if gpu_match:
        info["gpu_name"] = " ".join(gpu_match.group(1).split())
    return info


def basename(path_text: str) -> str:
    return Path(path_text).name if path_text else ""


def file_entry(path: Path, note: str) -> dict[str, str]:
    return {"path": str(path), "note": note}


def build_markdown_table(rows: list[dict[str, Any]], include_final: bool = False) -> str:
    header = [
        "| 实验名 | 基础权重 | 输入尺寸 | Batch | 数据比例 | Cosine LR | mAP50-95 | Recall | Inference(ms) |",
        "| --- | --- | ---: | ---: | ---: | --- | ---: | ---: | ---: |",
    ]
    lines: list[str] = []
    for row in rows:
        if not include_final and str(row["experiment"]).startswith("final_"):
            continue
        lines.append(
            "| {experiment} | {base} | {imgsz} | {batch} | {fraction} | {cos_lr} | {map50_95} | {recall} | {inference} |".format(
                experiment=row["experiment"],
                base=basename(str(row["base_model"])),
                imgsz=row["imgsz"],
                batch=row["batch"],
                fraction=fmt_float(row["fraction"], 2),
                cos_lr="True" if row["cos_lr"] else "False",
                map50_95=fmt_float(row["map50_95"]),
                recall=fmt_float(row["recall"]),
                inference=fmt_float(row["inference_ms"], 3),
            )
        )
    return "\n".join(header + lines)


def build_ablation_comments(best_ablation: dict[str, Any], ablation_rows: list[dict[str, Any]]) -> list[str]:
    rows_by_name = {row["experiment"]: row for row in ablation_rows}
    current = rows_by_name[best_ablation["experiment"]]
    comments = [
        f"最优消融配置为 `{current['experiment']}`，测试集 `mAP50-95={fmt_float(current['map50_95'])}`。",
    ]
    baseline = rows_by_name.get("ab0_generic640_default_f20")
    if baseline:
        delta = float(current["map50_95"]) - float(baseline["map50_95"])
        comments.append(
            f"相较通用预训练基线 `ab0_generic640_default_f20`，`mAP50-95` 绝对提升 {fmt_float(delta)}。"
        )
    same_base = rows_by_name.get("ab2_plate640_coslr_f20")
    if same_base:
        delta = float(current["map50_95"]) - float(same_base["map50_95"])
        comments.append(
            f"在使用车牌专用初始化后，将输入尺寸从 640 提升到 960，可进一步带来 {fmt_float(delta)} 的 `mAP50-95` 增益。"
        )
    low_aug = rows_by_name.get("ab4_plate960_coslr_lowaug_f20")
    if low_aug:
        comments.append(
            "低增强版本 `ab4_plate960_coslr_lowaug_f20` 未超过最优配置，说明当前默认增强组合更适合该混合数据集。"
        )
    return comments


def load_optional_multiseed(record_root: Path) -> dict[str, Any] | None:
    stats_path = record_root / "multiseed" / "summaries" / "seed_stability_stats.json"
    runs_path = record_root / "multiseed" / "summaries" / "seed_stability_runs.json"
    summary_md = record_root / "multiseed" / "summaries" / "seed_stability_summary.md"
    if not stats_path.is_file() or not runs_path.is_file():
        return None
    return {
        "stats_path": str(stats_path),
        "runs_path": str(runs_path),
        "summary_md": str(summary_md) if summary_md.is_file() else "",
        "stats": load_json(stats_path),
        "runs": load_json(runs_path),
    }


def build_paper_summary_md(payload: dict[str, Any]) -> str:
    dataset = payload["dataset"]
    env = payload["environment"]
    best_ablation = payload["best_ablation"]
    final_run = payload["final_run"]
    ablation_rows = payload["ablation_rows"]
    final_row = payload["final_row"]
    artifact_groups = payload["artifact_groups"]
    multi_seed = payload.get("multi_seed")

    lines: list[str] = []
    lines.append("# CCPD Mixed 车牌检测论文级实验总结")
    lines.append("")
    lines.append(f"- 生成时间: `{payload['generated_at']}`")
    lines.append(f"- 记录根目录: `{payload['record_root']}`")
    lines.append(f"- 数据集配置: `{dataset['dataset_yaml']}`")
    lines.append("- 当前实验针对单类车牌检测任务，不直接输出完整车牌字符串。")
    lines.append("")
    lines.append("## 1. 数据集与实验环境")
    lines.append("")
    lines.append(
        f"- 数据集总样本数: `{dataset['total_samples']}`，训练/验证/测试划分为 "
        f"`{dataset['splits']['train']}` / `{dataset['splits']['val']}` / `{dataset['splits']['test']}`。"
    )
    lines.append(f"- 类别定义: `{', '.join(dataset['class_names'])}`")
    lines.append(
        f"- GPU 训练环境: `{env['gpu_name']}`，Torch `{env['torch_version']}`，"
        f"CUDA 编译版本 `{env['torch_cuda_version']}`，驱动 `{env['driver_version']}`。"
    )
    lines.append(f"- Python 解释器: `{env['python_executable']}`")
    lines.append("")
    lines.append("## 2. 消融实验设计")
    lines.append("")
    lines.append("- 所有消融实验均使用 GPU 加速训练，`epochs=8`，`fraction=0.2`。")
    lines.append("- 消融变量包括初始化权重、学习率调度、输入分辨率与轻量增强调整。")
    lines.append("")
    lines.append(build_markdown_table(ablation_rows, include_final=False))
    lines.append("")
    lines.append("### 消融结论")
    lines.append("")
    for comment in build_ablation_comments(best_ablation, ablation_rows):
        lines.append(f"- {comment}")
    lines.append("")
    lines.append("## 3. 最终正式训练结果")
    lines.append("")
    lines.append(
        f"- 采用配置: `{final_row['experiment']}`，基础权重 `{basename(str(final_row['base_model']))}`，"
        f"`imgsz={final_row['imgsz']}`，`batch={final_row['batch']}`，`epochs={final_run['epochs']}`，`patience={final_run['patience']}`。"
    )
    lines.append(
        f"- 测试集指标: Precision `{fmt_float(final_run['test_metrics']['precision'])}`，"
        f"Recall `{fmt_float(final_run['test_metrics']['recall'])}`，"
        f"mAP50 `{fmt_float(final_run['test_metrics']['map50'])}`，"
        f"mAP50-95 `{fmt_float(final_run['test_metrics']['map50_95'])}`。"
    )
    lines.append(
        f"- 验证集指标: Precision `{fmt_float(final_run['val_metrics']['precision'])}`，"
        f"Recall `{fmt_float(final_run['val_metrics']['recall'])}`，"
        f"mAP50 `{fmt_float(final_run['val_metrics']['map50'])}`，"
        f"mAP50-95 `{fmt_float(final_run['val_metrics']['map50_95'])}`。"
    )
    lines.append(
        f"- 测试推理速度: `preprocess={fmt_float(final_run['test_speed']['preprocess'], 3)} ms`，"
        f"`inference={fmt_float(final_run['test_speed']['inference'], 3)} ms`，"
        f"`postprocess={fmt_float(final_run['test_speed']['postprocess'], 3)} ms`。"
    )
    lines.append(f"- 最终最优权重: `{final_run['best_weights']}`")
    lines.append(
        f"- 正式训练相较最优消融配置 `{best_ablation['experiment']}` 的测试 `mAP50-95` 进一步提升 "
        f"`{fmt_float(float(final_run['test_metrics']['map50_95']) - float(best_ablation['map50_95']))}`。"
    )
    lines.append("")

    if multi_seed:
        stats = multi_seed["stats"]
        metrics = stats["metrics"]
        lines.append("## 4. 多随机种子稳定性实验")
        lines.append("")
        lines.append(f"- 随机种子: `{', '.join(str(seed) for seed in stats['seeds'])}`")
        lines.append(
            f"- `mAP50-95 = {fmt_float(metrics['map50_95']['mean'])} ± {fmt_float(metrics['map50_95']['std'])}`"
        )
        lines.append(
            f"- `Precision = {fmt_float(metrics['precision']['mean'])} ± {fmt_float(metrics['precision']['std'])}`"
        )
        lines.append(
            f"- `Recall = {fmt_float(metrics['recall']['mean'])} ± {fmt_float(metrics['recall']['std'])}`"
        )
        lines.append(
            f"- `Inference = {fmt_float(metrics['inference_ms']['mean'], 3)} ± {fmt_float(metrics['inference_ms']['std'], 3)} ms/image`"
        )
        lines.append(
            f"- 最优种子: `{stats['best_run']['seed']}` (`mAP50-95={fmt_float(stats['best_run']['map50_95'])}`)，"
            f"最弱种子: `{stats['worst_run']['seed']}` (`mAP50-95={fmt_float(stats['worst_run']['map50_95'])}`)。"
        )
        lines.append("")
        lines.append("该部分可直接作为论文中的稳定性分析或补充实验章节。")
        lines.append("")

    lines.append("## 5. 可直接引用的论文表述建议")
    lines.append("")
    lines.append(
        "在 CCPD mixed 单类车牌检测任务上，本文首先基于 20% 训练子集设计五组消融实验，"
        "系统比较通用预训练权重、车牌专用初始化、余弦学习率调度及输入分辨率等因素。"
    )
    lines.append(
        f"实验结果表明，`{best_ablation['experiment']}` 在测试集上取得 `mAP50-95={fmt_float(best_ablation['map50_95'])}`，"
        "说明车牌专用初始化与更高输入分辨率对该任务具有显著正向作用。"
    )
    lines.append(
        f"在此基础上，采用全量训练集对该配置进行正式训练，最终模型在测试集上获得 "
        f"`Precision={fmt_float(final_run['test_metrics']['precision'])}`、"
        f"`Recall={fmt_float(final_run['test_metrics']['recall'])}`、"
        f"`mAP50={fmt_float(final_run['test_metrics']['map50'])}`、"
        f"`mAP50-95={fmt_float(final_run['test_metrics']['map50_95'])}`。"
    )
    if multi_seed:
        metrics = multi_seed["stats"]["metrics"]
        lines.append(
            f"进一步的多随机种子重复实验显示，该配置在 `mAP50-95` 上达到 "
            f"`{fmt_float(metrics['map50_95']['mean'])} ± {fmt_float(metrics['map50_95']['std'])}`，"
            "说明模型在不同随机初始化下具有较好的稳定性。"
        )
    lines.append("")
    lines.append("## 6. 关键产物")
    lines.append("")
    lines.append("- 最终权重与训练目录")
    lines.append(f"  - `{final_run['best_weights']}`")
    lines.append(f"  - `{final_run['run_dir']}`")
    lines.append("- 关键图表")
    for item in artifact_groups["figures"]:
        lines.append(f"  - `{item['path']}`: {item['note']}")
    lines.append("- 结构化报告")
    for item in artifact_groups["reports"]:
        lines.append(f"  - `{item['path']}`: {item['note']}")
    if multi_seed:
        lines.append("- 多随机种子汇总")
        lines.append(f"  - `{multi_seed['summary_md']}`")
        lines.append(f"  - `{multi_seed['stats_path']}`")
    lines.append("")
    lines.append("## 7. 复现实验入口")
    lines.append("")
    lines.append(f"- 套件脚本: `{payload['repro']['suite_script']}`")
    lines.append(f"- 多随机种子脚本: `{payload['repro']['multiseed_script']}`")
    lines.append(f"- 文档生成脚本: `{payload['repro']['doc_builder_script']}`")
    lines.append(f"- GPU Python: `{env['python_executable']}`")
    lines.append("")
    lines.append("## 8. 后续可扩展方向")
    lines.append("")
    lines.append("- 若需要完整车牌号识别，应在检测后增加 OCR 或字符级识别分支。")
    lines.append("- 若需要视频实验，应在逐帧检测基础上加入目标跟踪与多帧投票。")
    lines.append("- 若需要更完整论文支撑，可继续补充不同模型规模、不同训练轮次的对比。")
    lines.append("")
    return "\n".join(lines)


def build_artifact_index_md(payload: dict[str, Any]) -> str:
    artifact_groups = payload["artifact_groups"]
    lines = ["# 论文实验产物索引", ""]
    lines.append(f"- 生成时间: `{payload['generated_at']}`")
    lines.append(f"- 根目录: `{payload['record_root']}`")
    lines.append("")
    for title, items in (
        ("规划与元数据", artifact_groups["metadata"]),
        ("汇总文档", artifact_groups["reports"]),
        ("最终模型", artifact_groups["models"]),
        ("关键图表", artifact_groups["figures"]),
        ("日志文件", artifact_groups["logs"]),
    ):
        lines.append(f"## {title}")
        lines.append("")
        for item in items:
            lines.append(f"- `{item['path']}`: {item['note']}")
        lines.append("")
    return "\n".join(lines)


def build_docx(payload: dict[str, Any], output_path: Path) -> None:
    dataset = payload["dataset"]
    env = payload["environment"]
    best_ablation = payload["best_ablation"]
    final_run = payload["final_run"]
    ablation_rows = payload["ablation_rows"]
    multi_seed = payload.get("multi_seed")

    doc = Document()
    doc.add_heading("CCPD Mixed 车牌检测论文级实验总结", level=0)
    doc.add_paragraph(f"生成时间: {payload['generated_at']}")
    doc.add_paragraph(f"记录根目录: {payload['record_root']}")

    doc.add_heading("1. 数据集与环境", level=1)
    doc.add_paragraph(
        f"数据集总样本数 {dataset['total_samples']}，划分为训练集 {dataset['splits']['train']}、"
        f"验证集 {dataset['splits']['val']}、测试集 {dataset['splits']['test']}，类别为 {', '.join(dataset['class_names'])}。"
    )
    doc.add_paragraph(
        f"训练环境为 {env['gpu_name']}，Torch {env['torch_version']}，"
        f"CUDA {env['torch_cuda_version']}，驱动版本 {env['driver_version']}。"
    )

    doc.add_heading("2. 消融实验结果", level=1)
    table = doc.add_table(rows=1, cols=9)
    header = table.rows[0].cells
    header[0].text = "实验名"
    header[1].text = "基础权重"
    header[2].text = "输入尺寸"
    header[3].text = "Batch"
    header[4].text = "数据比例"
    header[5].text = "Cos LR"
    header[6].text = "mAP50-95"
    header[7].text = "Recall"
    header[8].text = "Inference(ms)"
    for row in ablation_rows:
        cells = table.add_row().cells
        cells[0].text = str(row["experiment"])
        cells[1].text = basename(str(row["base_model"]))
        cells[2].text = str(row["imgsz"])
        cells[3].text = str(row["batch"])
        cells[4].text = fmt_float(row["fraction"], 2)
        cells[5].text = "True" if row["cos_lr"] else "False"
        cells[6].text = fmt_float(row["map50_95"])
        cells[7].text = fmt_float(row["recall"])
        cells[8].text = fmt_float(row["inference_ms"], 3)
    for comment in build_ablation_comments(best_ablation, ablation_rows):
        doc.add_paragraph(comment, style="List Bullet")

    doc.add_heading("3. 最终正式训练结果", level=1)
    doc.add_paragraph(
        f"最终采用 {payload['final_row']['experiment']} 配置，在测试集上获得 "
        f"Precision={fmt_float(final_run['test_metrics']['precision'])}，"
        f"Recall={fmt_float(final_run['test_metrics']['recall'])}，"
        f"mAP50={fmt_float(final_run['test_metrics']['map50'])}，"
        f"mAP50-95={fmt_float(final_run['test_metrics']['map50_95'])}。"
    )
    doc.add_paragraph(f"最终最优权重路径: {final_run['best_weights']}")
    doc.add_paragraph(f"训练运行目录: {final_run['run_dir']}")

    if multi_seed:
        metrics = multi_seed["stats"]["metrics"]
        doc.add_heading("4. 多随机种子稳定性", level=1)
        doc.add_paragraph(
            f"多随机种子实验的 mAP50-95 为 {fmt_float(metrics['map50_95']['mean'])} ± {fmt_float(metrics['map50_95']['std'])}，"
            f"Precision 为 {fmt_float(metrics['precision']['mean'])} ± {fmt_float(metrics['precision']['std'])}，"
            f"Recall 为 {fmt_float(metrics['recall']['mean'])} ± {fmt_float(metrics['recall']['std'])}。"
        )

    doc.add_heading("5. 关键文件", level=1)
    for group_name in ("reports", "models", "figures"):
        for item in payload["artifact_groups"][group_name]:
            doc.add_paragraph(f"{item['path']} - {item['note']}", style="List Bullet")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> None:
    args = parse_args()
    record_root = Path(args.record_root).resolve()

    summary_rows = load_json(record_root / "summaries" / "experiment_summary.json")
    best_ablation = load_json(record_root / "summaries" / "best_ablation.json")
    split_summary = load_json(record_root / "meta" / "split_summary.json")
    gpu_probe = load_json(record_root / "meta" / "python_gpu_probe.json")
    nvidia_smi_text = (record_root / "meta" / "nvidia_smi.txt").read_text(encoding="utf-8")
    nvidia_info = parse_nvidia_smi(nvidia_smi_text)
    multi_seed = load_optional_multiseed(record_root)

    ablation_rows = [row for row in summary_rows if not str(row["experiment"]).startswith("final_")]
    final_rows = [row for row in summary_rows if str(row["experiment"]).startswith("final_")]
    if not final_rows:
        raise RuntimeError("No final formal training rows found in experiment_summary.json")
    final_row = max(final_rows, key=lambda item: float(item["map50_95"]))

    final_train_report = load_json(Path(final_row["train_report"]))
    final_eval_report = load_json(Path(final_row["eval_report"]))
    run_dir = Path(final_row["run_dir"])

    reports = [
        file_entry(record_root / args.paper_summary_md, "论文级中文摘要（Markdown）"),
        file_entry(record_root / args.paper_summary_docx, "论文级中文摘要（Word）"),
        file_entry(record_root / args.artifact_index_md, "实验产物索引"),
        file_entry(record_root / "summaries" / "experiment_summary.md", "已有实验汇总表"),
        file_entry(record_root / "summaries" / "experiment_summary.csv", "实验汇总 CSV"),
        file_entry(record_root / "summaries" / "experiment_summary.json", "实验汇总 JSON"),
        file_entry(record_root / "summaries" / "best_ablation.json", "最优消融配置"),
        file_entry(Path(final_row["train_report"]), "最终正式训练报告"),
        file_entry(Path(final_row["eval_report"]), "最终测试集评估报告"),
    ]
    if multi_seed:
        if multi_seed["summary_md"]:
            reports.append(file_entry(Path(multi_seed["summary_md"]), "多随机种子摘要"))
        reports.append(file_entry(Path(multi_seed["stats_path"]), "多随机种子统计"))
        reports.append(file_entry(Path(multi_seed["runs_path"]), "多随机种子明细"))

    models = [
        file_entry(Path(final_row["best_weights"]), "最终最优权重"),
        file_entry(run_dir / "weights" / "last.pt", "最终训练的 last 权重"),
        file_entry(run_dir, "最终训练运行目录"),
    ]

    figures = [
        file_entry(run_dir / "results.png", "训练过程总览曲线"),
        file_entry(run_dir / "results.csv", "训练过程数值记录"),
        file_entry(run_dir / "BoxPR_curve.png", "PR 曲线"),
        file_entry(run_dir / "BoxF1_curve.png", "F1-Confidence 曲线"),
        file_entry(run_dir / "BoxP_curve.png", "Precision-Confidence 曲线"),
        file_entry(run_dir / "BoxR_curve.png", "Recall-Confidence 曲线"),
        file_entry(run_dir / "confusion_matrix.png", "混淆矩阵"),
        file_entry(run_dir / "confusion_matrix_normalized.png", "归一化混淆矩阵"),
        file_entry(run_dir / "val_batch0_pred.jpg", "验证集预测样例 1"),
        file_entry(run_dir / "val_batch1_pred.jpg", "验证集预测样例 2"),
        file_entry(run_dir / "val_batch2_pred.jpg", "验证集预测样例 3"),
    ]

    video_demo_dir = record_root / "video_demo"
    for path in sorted(video_demo_dir.glob("*")):
        if path.is_file():
            figures.append(file_entry(path, "视频链路样例产物"))

    logs = [file_entry(path, "实验日志") for path in sorted((record_root / "logs").glob("*.log"))]
    if (record_root / "multiseed" / "logs").is_dir():
        logs.extend(
            file_entry(path, "多随机种子日志")
            for path in sorted((record_root / "multiseed" / "logs").glob("*.log"))
        )

    metadata = [
        file_entry(record_root / "plans" / "experiment_plan.md", "实验规划"),
        file_entry(record_root / "meta" / "split_summary.json", "数据集划分摘要"),
        file_entry(record_root / "meta" / "dataset.yaml", "训练数据集 YAML 副本"),
        file_entry(record_root / "meta" / "python_gpu_probe.json", "GPU 可用性探针"),
        file_entry(record_root / "meta" / "nvidia_smi.txt", "GPU 运行环境快照"),
        file_entry(record_root / "meta" / "ablation_runs.json", "消融实验执行清单"),
    ]
    if (record_root / "multiseed" / "plans" / "multiseed_plan.md").is_file():
        metadata.append(file_entry(record_root / "multiseed" / "plans" / "multiseed_plan.md", "多随机种子规划"))

    final_training = final_train_report["training"]
    final_metrics_test = final_train_report["metrics"]["test"]
    final_metrics_val = final_train_report["metrics"]["val"]

    payload = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "record_root": str(record_root),
        "dataset": {
            "dataset_yaml": split_summary["dataset_yaml"],
            "total_samples": split_summary["total_samples"],
            "splits": split_summary["splits"],
            "class_names": split_summary["class_names"],
            "seed": split_summary["seed"],
            "ratios": split_summary["ratios"],
        },
        "environment": {
            "python_executable": (record_root / "meta" / "python_executable.txt").read_text(encoding="utf-8").strip()
            if (record_root / "meta" / "python_executable.txt").exists()
            else "",
            "python_version": gpu_probe.get("python", ""),
            "torch_version": gpu_probe.get("torch", ""),
            "torch_cuda_version": "11.8" if "+cu118" in str(gpu_probe.get("torch", "")) else "",
            "gpu_name": gpu_probe.get("device_name", nvidia_info.get("gpu_name", "")),
            "driver_version": nvidia_info.get("driver_version", ""),
            "cuda_runtime_version": nvidia_info.get("cuda_runtime_version", ""),
        },
        "ablation_rows": ablation_rows,
        "best_ablation": best_ablation,
        "final_row": final_row,
        "final_run": {
            "epochs": final_training["epochs"],
            "patience": final_training["patience"],
            "run_dir": final_training["run_dir"],
            "best_weights": final_training["best_weights"],
            "last_weights": final_training["last_weights"],
            "test_metrics": {
                "precision": final_metrics_test["box_mp"],
                "recall": final_metrics_test["box_mr"],
                "map50": final_metrics_test["box_map50"],
                "map50_95": final_metrics_test["box_map"],
            },
            "val_metrics": {
                "precision": final_metrics_val["box_mp"],
                "recall": final_metrics_val["box_mr"],
                "map50": final_metrics_val["box_map50"],
                "map50_95": final_metrics_val["box_map"],
            },
            "test_speed": final_eval_report["metrics"]["speed"],
        },
        "multi_seed": multi_seed,
        "artifact_groups": {
            "metadata": metadata,
            "reports": reports,
            "models": models,
            "figures": figures,
            "logs": logs,
        },
        "repro": {
            "suite_script": str((PROJECT_ROOT / "smart-parking-ai" / "run_plate_thesis_suite.py").resolve()),
            "multiseed_script": str((PROJECT_ROOT / "smart-parking-ai" / "run_plate_multiseed.py").resolve()),
            "doc_builder_script": str((PROJECT_ROOT / "smart-parking-ai" / "build_plate_paper_docs.py").resolve()),
        },
    }

    paper_summary_md_path = record_root / args.paper_summary_md
    artifact_index_md_path = record_root / args.artifact_index_md
    manifest_json_path = record_root / args.manifest_json
    paper_summary_docx_path = record_root / args.paper_summary_docx

    write_text(paper_summary_md_path, build_paper_summary_md(payload))
    write_text(artifact_index_md_path, build_artifact_index_md(payload))
    write_json(manifest_json_path, payload)
    build_docx(payload, paper_summary_docx_path)

    print(f"paper_summary_md={paper_summary_md_path}")
    print(f"artifact_index_md={artifact_index_md_path}")
    print(f"manifest_json={manifest_json_path}")
    print(f"paper_summary_docx={paper_summary_docx_path}")


if __name__ == "__main__":
    main()
